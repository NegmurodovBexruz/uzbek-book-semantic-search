import os
import re
from docx import Document

ARABIC_RE = re.compile(
    r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF"
    r"\uFB50-\uFDFF\uFE70-\uFEFF]+"
)

def remove_arabic(text: str) -> str:
    return ARABIC_RE.sub("", text)

def normalize_space(text: str) -> str:
    text = text.replace("\u00A0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

def read_docx_with_headings_clean(path: str):
    doc = Document(path)
    current_heading = ""
    out = []

    for p in doc.paragraphs:
        raw = (p.text or "").strip()
        if not raw:
            continue

        style = (p.style.name or "").lower() if p.style else ""
        is_heading = "heading" in style

        cleaned = normalize_space(remove_arabic(raw))
        if not cleaned:
            if is_heading:
                current_heading = ""
            continue

        if is_heading:
            current_heading = cleaned
            continue

        out.append((cleaned, current_heading))

    return out

def is_probably_heading(p, cleaned_text: str) -> bool:
    """
    DOCX paragraf sarlavhami yo'qmi - heuristika.
    """
    # 1) Word style Heading
    style = (p.style.name or "").lower() if p.style else ""
    if "heading" in style:
        return True

    # 2) Juda uzun bo'lmasa + oxiri nuqta bilan tugamasa (sarlavha odatda nuqta bilan tugamaydi)
    if len(cleaned_text) > 120:
        return False

    # 3) Raqamli bo'limlar: 1. , 1.1 , 2.3.4 ...
    if re.match(r"^\d+(\.\d+)*[\)\.]?\s+\S+", cleaned_text):
        return True


    upper = cleaned_text.upper()
    if upper.startswith("BOB") or upper.startswith("FASL") or upper.startswith("MAVZU"):
        return True


    runs = getattr(p, "runs", [])
    if runs:
        total_chars = sum(len(r.text or "") for r in runs)
        bold_chars = sum(len(r.text or "") for r in runs if r.bold)
        if total_chars > 0 and (bold_chars / total_chars) >= 0.6:
            return True


    letters = re.sub(r"[^A-Za-zА-Яа-яЎўҚқҒғҲҳ]", "", cleaned_text)
    if letters and letters.isupper() and len(letters) >= 8:
        return True

    return False


def read_docx_sections_clean(path: str):
    doc = Document(path)
    sections = []
    current_heading = "NO_HEADING"
    current_lines = []

    def flush():
        nonlocal current_lines, current_heading
        text = "\n".join(current_lines).strip()
        if text:
            sections.append({"heading": current_heading, "text": text})
        current_lines = []

    for p in doc.paragraphs:
        raw = (p.text or "").strip()
        if not raw:
            continue

        cleaned = normalize_space(remove_arabic(raw))
        if not cleaned:
            continue

        if is_probably_heading(p, cleaned):
            flush()
            current_heading = cleaned
            continue

        current_lines.append(cleaned)

    flush()
    return sections

def book_title_from_path(path: str) -> str:
    return os.path.splitext(os.path.basename(path))[0]
